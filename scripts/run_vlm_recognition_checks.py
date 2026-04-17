#!/usr/bin/env python3
"""Проверка корректности распознавания изображений через VLM в runtime и ingest.

Сценарий:
1) Генерирует тестовые изображения разных форматов (png, jpeg, bmp, tiff).
2) Runtime-check: прогоняет VisionService.analyze_attachments() в режиме VLM.
3) Ingest-check: встраивает изображения в PDF/DOCX, извлекает image_assets парсерами,
   затем прогоняет VisionService.build_document_image_chunks() в режиме VLM.
"""

from __future__ import annotations

import argparse
import os
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path


REPO_ROOT = Path(__file__).resolve().parents[1]
APP_ROOT = REPO_ROOT / 'app'
if str(APP_ROOT) not in sys.path:
    sys.path.insert(0, str(APP_ROOT))

from docx import Document  # noqa: E402
from docx.shared import Inches  # noqa: E402
from PIL import Image, ImageDraw, ImageFont  # noqa: E402

from src.api.schemas import AttachmentItem  # noqa: E402
from src.core.settings import settings  # noqa: E402
from src.ingest.parsers.docx_parser import parse_docx  # noqa: E402
from src.ingest.parsers.pdf_parser import parse_pdf  # noqa: E402
from src.vision.service import VisionService  # noqa: E402


@dataclass
class CheckResult:
    name: str
    ok: bool
    details: str


def _make_test_image(path: Path, *, label: str, marker: str) -> None:
    img = Image.new('RGB', (1280, 720), color=(247, 250, 252))
    draw = ImageDraw.Draw(img)
    font = ImageFont.load_default()
    draw.rectangle((0, 0, 1280, 80), fill=(30, 41, 59))
    draw.text((24, 30), f'VLM CHECK: {label}', fill=(255, 255, 255), font=font)
    draw.text((40, 150), f'Error code: HTTP 503', fill=(15, 23, 42), font=font)
    draw.text((40, 200), f'Access denied while opening report', fill=(15, 23, 42), font=font)
    draw.text((40, 250), f'Marker: {marker}', fill=(15, 23, 42), font=font)
    img.save(path)


def _generate_runtime_images(work_dir: Path, marker: str) -> dict[str, Path]:
    runtime_dir = work_dir / 'runtime'
    runtime_dir.mkdir(parents=True, exist_ok=True)
    files = {
        'png': runtime_dir / 'runtime_case.png',
        'jpeg': runtime_dir / 'runtime_case.jpeg',
        'bmp': runtime_dir / 'runtime_case.bmp',
        'tiff': runtime_dir / 'runtime_case.tiff',
    }
    for fmt, path in files.items():
        _make_test_image(path, label=fmt.upper(), marker=marker)
    return files


def _build_docx_with_images(path: Path, images: list[Path]) -> None:
    doc = Document()
    doc.add_heading('VLM ingest check', level=1)
    doc.add_paragraph('Документ для проверки извлечения и распознавания изображений.')
    for image in images:
        doc.add_paragraph(f'Image asset: {image.name}')
        doc.add_picture(str(image), width=Inches(6.8))
    doc.save(path)


def _build_pdf_with_image(path: Path, image_path: Path) -> None:
    image = Image.open(image_path).convert('RGB')
    image.save(path, 'PDF', resolution=150.0)


def _runtime_check(images: dict[str, Path], *, question: str) -> list[CheckResult]:
    service = VisionService()
    checks: list[CheckResult] = []
    for fmt, image_path in images.items():
        evidence = service.analyze_attachments([AttachmentItem(image_path=str(image_path))], question=question)
        ok = bool(
            evidence
            and evidence[0].summary
            and 'Метод: VLM' in evidence[0].summary
            and evidence[0].confidence > 0.2
        )
        details = (
            f"summary={evidence[0].summary[:120]!r}, confidence={evidence[0].confidence:.2f}, "
            f"ocr_text_len={len(evidence[0].ocr_text)}"
            if evidence
            else 'no visual evidence returned'
        )
        checks.append(CheckResult(name=f'runtime {fmt}', ok=ok, details=details))
    return checks


def _ingest_check(work_dir: Path, images: dict[str, Path]) -> list[CheckResult]:
    ingest_dir = work_dir / 'ingest'
    ingest_dir.mkdir(parents=True, exist_ok=True)
    docx_path = ingest_dir / 'vlm_ingest_case.docx'
    pdf_path = ingest_dir / 'vlm_ingest_case.pdf'

    _build_docx_with_images(docx_path, [images['png'], images['jpeg'], images['bmp'], images['tiff']])
    _build_pdf_with_image(pdf_path, images['png'])

    parsed_docx = parse_docx(str(docx_path), source_type='csv_ans_docs', doc_id='vlm_docx_case')
    parsed_pdf = parse_pdf(str(pdf_path), source_type='csv_ans_docs', doc_id='vlm_pdf_case')

    service = VisionService()
    docx_chunks = service.build_document_image_chunks(
        parsed_docx.get('image_assets', []),
        doc_id='vlm_docx_case',
        source_type='csv_ans_docs',
    )
    pdf_chunks = service.build_document_image_chunks(
        parsed_pdf.get('image_assets', []),
        doc_id='vlm_pdf_case',
        source_type='csv_ans_docs',
    )

    def _has_non_empty_vlm_payload(chunks: list[dict]) -> bool:
        for chunk in chunks:
            text = str(chunk.get('text', ''))
            if 'VLM:\n' not in text:
                return False
            payload = text.split('VLM:\n', maxsplit=1)[1].strip()
            if len(payload) < 10:
                return False
        return bool(chunks)

    checks = [
        CheckResult(
            name='ingest docx extraction+vlm',
            ok=_has_non_empty_vlm_payload(docx_chunks),
            details=f'docx_assets={len(parsed_docx.get("image_assets", []))}, docx_chunks={len(docx_chunks)}',
        ),
        CheckResult(
            name='ingest pdf extraction+vlm',
            ok=_has_non_empty_vlm_payload(pdf_chunks),
            details=f'pdf_assets={len(parsed_pdf.get("image_assets", []))}, pdf_chunks={len(pdf_chunks)}',
        ),
    ]
    return checks


def main() -> int:
    parser = argparse.ArgumentParser(description='Проверка VLM распознавания изображений (runtime+ingest).')
    parser.add_argument('--work-dir', default='data/vision_vlm_checks', help='Рабочий каталог для test assets.')
    parser.add_argument('--marker', default='VLM-CHECK-503', help='Маркер, добавляемый в тестовые изображения.')
    parser.add_argument(
        '--question',
        default='Опиши, что видно на скриншоте и есть ли признаки ошибки.',
        help='Вопрос для runtime-проверки.',
    )
    parser.add_argument(
        '--keep-assets',
        action='store_true',
        help='Не удалять созданные test assets после завершения.',
    )
    args = parser.parse_args()

    os.environ['VISION_RUNTIME_MODE'] = 'vlm'
    os.environ['VISION_INGEST_MODE'] = 'vlm'
    settings.vision_runtime_mode = 'vlm'
    settings.vision_ingest_mode = 'vlm'

    work_dir = (REPO_ROOT / args.work_dir).resolve()
    work_dir.mkdir(parents=True, exist_ok=True)

    images = _generate_runtime_images(work_dir, marker=args.marker)
    checks: list[CheckResult] = []
    checks.extend(_runtime_check(images, question=args.question))
    checks.extend(_ingest_check(work_dir, images))

    print('=== VLM recognition checks ===')
    for item in checks:
        status = 'PASS' if item.ok else 'FAIL'
        print(f'[{status}] {item.name}: {item.details}')

    ok_all = all(item.ok for item in checks)
    if ok_all:
        print('RESULT: PASS')
    else:
        print('RESULT: FAIL')

    if not args.keep_assets:
        shutil.rmtree(work_dir, ignore_errors=True)
    return 0 if ok_all else 1


if __name__ == '__main__':
    raise SystemExit(main())
